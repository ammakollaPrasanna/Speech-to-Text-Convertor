from flask import Flask, request, jsonify, send_file
from flask_cors import CORS # Added for cross-origin requests
from docx import Document
from docx.shared import Inches # Added, though not used in the provided snippet
import io
from datetime import datetime # Retained for timestamping

import os # Retained for file path operations


# Install these libraries:
# pip install Flask python-docx SpeechRecognition pydub nltk Flask-CORS
# You might also need to install ffmpeg (https://ffmpeg.org/download.html)
# and ensure its path is in your system's environment variables if you plan
# to process audio files like .mp3, .wav etc. with pydub.

app = Flask(_name_)
CORS(app)  # Enable CORS for cross-origin requests

# --- Conceptual Backend Speech-to-Text (if receiving audio files) ---
# Note: The frontend is currently doing client-side speech-to-text.
# This section would be used if you wanted to upload audio files to the backend
# for transcription, e.g., for longer recordings or batch processing.
try:
    import speech_recognition as sr
    from pydub import AudioSegment
    AUDIO_PROCESSING_ENABLED = True
except ImportError:
    print("SpeechRecognition or pydub not found. Server-side audio processing disabled.")
    print("Install with: pip install SpeechRecognition pydub")
    print("Also ensure you have ffmpeg installed and in your PATH.")
    AUDIO_PROCESSING_ENABLED = False
except Exception as e:
    print(f"Error importing audio processing libraries: {e}")
    AUDIO_PROCESSING_ENABLED = False


def transcribe_audio_file(audio_path):
    """
    Transcribes an audio file using Google Web Speech API (online).
    For offline or more robust solutions, consider paid APIs (e.g., Google Cloud Speech-to-Text)
    or open-source models like Vosk or DeepSpeech.
    """
    if not AUDIO_PROCESSING_ENABLED:
        return None, "Audio processing libraries not available."

    r = sr.Recognizer()
    try:
        # Convert audio to WAV if it's not already
        if not audio_path.lower().endswith('.wav'):
            audio = AudioSegment.from_file(audio_path)
            wav_path = "temp_audio.wav"
            audio.export(wav_path, format="wav")
            audio_file = sr.AudioFile(wav_path)
        else:
            audio_file = sr.AudioFile(audio_path)

        with audio_file as source:
            # Adjusts for ambient noise
            r.adjust_for_ambient_noise(source)
            audio_data = r.record(source) # read the entire audio file

        # Recognize speech using Google Web Speech API
        text = r.recognize_google(audio_data)
        return text, None
    except sr.UnknownValueError:
        return None, "Google Speech Recognition could not understand audio"
    except sr.RequestError as e:
        return None, f"Could not request results from Google Speech Recognition service; {e}"
    except Exception as e:
        return None, f"An error occurred during audio transcription: {e}"
    finally:
        # Clean up temporary WAV file if created
        if 'wav_path' in locals() and os.path.exists(wav_path):
            os.remove(wav_path)


# --- NLTK Integration (Conceptual) ---
try:
    import nltk
    # Ensure you have the 'punkt' tokenizer downloaded for sentence tokenization
    # nltk.download('punkt') # Uncomment and run once if you haven't downloaded it
    NLTK_ENABLED = True
except ImportError:
    print("NLTK not found. NLTK processing disabled.")
    print("Install with: pip install nltk")
    NLTK_ENABLED = False

def process_text_with_nltk(text_content):
    """
    Applies a simple NLTK operation (sentence tokenization) to the text.
    You can expand this for summarization, keyword extraction, etc.
    This function is retained for potential future use but not directly
    called in the current /generate-word endpoint as per user's latest request.
    """
    if not NLTK_ENABLED:
        return text_content, "NLTK is not enabled for text processing."
    try:
        # Example: Sentence tokenization
        sentences = nltk.sent_tokenize(text_content)
        processed_text = "\n".join(sentences) # Join sentences with newlines
        return processed_text, None
    except Exception as e:
        return text_content, f"NLTK processing error: {e}"


# --- API Endpoint to Generate Word Document (User's new code) ---
@app.route('/generate-word', methods=['POST'])
def generate_word():
    try:
        data = request.json

        doc = Document()
        doc.add_heading('Minutes of Meeting', 0)

        # === Fixed Fields ===
        doc.add_heading('Meeting Information', level=1)
        doc.add_paragraph(f"Variant Name: {data.get('variantName', '')}")
        doc.add_paragraph(f"Part Name: {data.get('partName', '')}")
        doc.add_paragraph(f"Subject: {data.get('subject', '')}")
        doc.add_paragraph(f"Meeting Number: {data.get('meetingNumber', '')}")
        doc.add_paragraph(f"Title: {data.get('title', '')}")
        doc.add_paragraph(f"Keywords: {data.get('keywords', '')}")
        doc.add_paragraph(f"Meeting Date: {data.get('date', '')}")
        doc.add_paragraph(f"Meeting Day: {data.get('day', '')}")
        doc.add_paragraph(f"Meeting Time: {data.get('time', '')}")
        doc.add_paragraph(f"Venue: {data.get('venue', '')}")
        doc.add_paragraph(f"Members: {data.get('members', '')}")

        doc.add_paragraph()  # Blank line

        # === Helper function to add sections ===
        def add_section(title, items, field='text'):
            doc.add_heading(title, level=2)
            if not items:
                doc.add_paragraph("None")
            for i, item in enumerate(items, 1):
                # Apply NLTK processing if desired for these fields
                # processed_item_text, _ = process_text_with_nltk(item.get(field, ''))
                doc.add_paragraph(f"{i}. {item.get(field, '')}") # Using original text for now

        # Discussions
        add_section("Discussions", data.get('discussions', []))

        # Action Points (with multiple columns)
        action_points = data.get('actionPoints', [])
        doc.add_heading("Action Points and Presentations", level=2)
        if action_points:
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Action'
            hdr_cells[1].text = 'Responsibility'
            hdr_cells[2].text = 'Others'
            hdr_cells[3].text = 'PDC'
            hdr_cells[4].text = 'Work Center'
            for item in action_points:
                row = table.add_row().cells
                row[0].text = item.get('action', '')
                row[1].text = item.get('responsibility', '')
                row[2].text = item.get('others', '')
                row[3].text = item.get('pdc', '')
                row[4].text = item.get('workcenter', '')
        else:
            doc.add_paragraph("None")

        # Innovations
        add_section("Innovativeness / Lessons Learnt", data.get('innovations', []))

        # Decisions
        add_section("Decisions", data.get('decisions', []))

        # Add a footer with generation timestamp (retained from previous version)
        doc.add_paragraph(f'\nGenerated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')


        # Save to memory
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)

        # Determine filename based on new fields
        filename = f"MoM_{data.get('meetingNumber', 'Unknown')}_{data.get('date', 'Unknown')}.docx"
        filename = filename.replace('/', '-').replace(':', '') # Sanitize for filenames

        return send_file(
            mem_file,
            download_name=filename,
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# --- API Endpoint to Transcribe Uploaded Audio (Optional, retained from previous version) ---
@app.route('/api/transcribe_uploaded_audio', methods=['POST'])
def transcribe_uploaded_audio():
    """
    Receives an audio file, transcribes it using SpeechRecognition,
    and returns the transcribed text.
    """
    if 'audio_file' not in request.files:
        return jsonify({'success': False, 'message': 'No audio file part in the request'}), 400

    audio_file = request.files['audio_file']
    if audio_file.filename == '':
        return jsonify({'success': False, 'message': 'No selected audio file'}), 400

    if audio_file:
        # Save the uploaded file temporarily
        temp_audio_path = os.path.join('/tmp', audio_file.filename) # Use /tmp or a dedicated uploads folder
        audio_file.save(temp_audio_path)

        text, error = transcribe_audio_file(temp_audio_path)
        os.remove(temp_audio_path) # Clean up the temporary file

        if text:
            # Apply NLTK processing to the transcribed text if needed here
            # For this example, we'll return raw transcribed text.
            return jsonify({'success': True, 'transcribedText': text})
        else:
            return jsonify({'success': False, 'message': error}), 500
    return jsonify({'success': False, 'message': 'Unknown error'}), 500


if _name_ == '_main_':
    # You would typically run this Flask app using a production-ready WSGI server
    # like Gunicorn or uWSGI, not directly with app.run() in production.
    app.run(debug=True, port=5000) # Run on port 5000